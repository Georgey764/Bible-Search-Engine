import json, hashlib, os
from docx import Document

# book = "Genesis"
# chapter = 1
books = ["Genesis","Exodus","Leviticus","Numbers","Deuteronomy","Joshua","Judges","Ruth","1 Samuel","2 Samuel","1 Kings","2 Kings","1 Chronicles","2 Chronicles","Ezra","Nehemiah","Esther","Job","Psalms","Proverbs","Ecclesiastes","Song of Solomon","Isaiah","Jeremiah","Lamentations","Ezekiel","Daniel","Hosea","Joel","Amos","Obadiah","Jonah","Micah","Nahum","Habakkuk","Zephaniah","Haggai","Zechariah","Malachi","Matthew","Mark","Luke","John","Acts","Romans","1 Corinthians","2 Corinthians","Galatians","Ephesians","Philippians","Colossians","1 Thessalonians","2 Thessalonians","1 Timothy","2 Timothy","Titus","Philemon","Hebrews","James","1 Peter","2 Peter","1 John","2 John","3 John","Jude","Revelation"]
inverse_index = {}
# document_name = f"{book}_ch_{chapter}.docx"
# document = Document(f"./Books/{book}/{document_name}")

def make_name(term, max_len=4):
	query = term.lower().encode("UTF-8")
	return "{}.txt".format(hashlib.md5(query).hexdigest()[:max_len])

def make_json(keyword, inverse_index, current_document, starting_index):
	max_ngram = 6
	for i in range(3, min(max_ngram, len(keyword)) + 1):
		ngram = keyword[:i]
		if (ngram not in inverse_index) and len(ngram) != 0:
			file_name = make_name(ngram.lower())
			inverse_index[ngram] = file_name
			txt_json = {
				current_document: [starting_index]
			}
			with open(f"./hashed_documents/{file_name}", "a") as file:
				file.write(f"\n{ngram}\t{json.dumps(txt_json)}")
		elif len(ngram) != 0:
			file_name = inverse_index[ngram]
			txt_json = {}
			all_lines = []
			with open(f"./hashed_documents/{file_name}", "r") as file:
				all_lines = file.readlines()

			found = False
			for index, line in enumerate(all_lines):
				if line.split("\t")[0].lower() == ngram.lower():
					found = True
					txt_json = json.loads(line.split("\t")[1])
					if current_document not in txt_json:
						txt_json[current_document] = [starting_index]
					else:
						txt_json[current_document].append(starting_index)
					all_lines[index] = f"{ngram}\t{json.dumps(txt_json)}"

			if len(ngram) != 0 and found == False:
				txt_json = {
					current_document: [starting_index]
				}
				all_lines.append(f"{ngram}\t{json.dumps(txt_json)}")

			with open(f"./hashed_documents/{file_name}", "w") as file:
				file.write("\n".join([line.strip() for line in all_lines]))

def hash_book(book_name):
	num_of_chap = len(os.listdir(f"./Books/{book_name}")) - 1
	for i in range(1, num_of_chap + 1):
		print("book: " + book_name + "\nchapter: " + str(i))
		document_name = f"{book_name}_ch_{i}.docx"
		document = Document(f"./Books/{book_name}/{document_name}")
		text = ""
		stop_words = {"a","about","above","after","again","against","all","am","an","and","any","are","aren't","as","at","be","because","been","before","being","below","between","both","but","by","can't","cannot","could","couldn't","did","didn't","do","does","doesn't","doing","don't","down","during","each","few","for","from","further","had","hadn't","has","hasn't","have","haven't","having","he","he'd","he'll","he's","her","here","here's","hers","herself","him","himself","his","how","how's","i","i'd","i'll","i'm","i've","if","in","into","is","isn't","it","it's","its","itself","let's","me","more","most","mustn't","my","myself","no","nor","not","of","off","on","once","only","or","other","ought","our","ours","ourselves","out","over","own","same","shan't","she","she'd","she'll","she's","should","shouldn't","so","some","such","than","that","that's","the","their","theirs","them","themselves","then","there","there's","these","they","they'd","they'll","they're","they've","this","those","through","to","too","under","until","up","very","was","wasn't","we","we'd","we'll","we're","we've","were","weren't","what","what's","when","when's","where","where's","which","while","who","who's","whom","why","why's","with","won't","would","wouldn't","you","you'd","you'll","you're","you've","your","yours","yourself","yourselves"}

		for paragraph in document.paragraphs:
			text = text + " " + paragraph.text
		
		for line_index, line in enumerate(text.split("\n")):
			for word_index, word in enumerate(line.split(" ")):
				if not word.lower() in stop_words and word != "":
					starting_index = 0
					for i in range(line_index):
						starting_index += len(text.split("\n")[i])
					starting_index += word_index
					word = word.strip(" \t\n0123456789.:[(<>)].,'\"‘’?--•…").lower()
					make_json(word, inverse_index, document_name, starting_index)

def main():
	for book in books:
		hash_book(book)

if __name__ == "__main__":
	main()

# import json, hashlib, os
# from docx import Document

# # book = "Genesis"
# # chapter = 1
# books = ["Genesis","Exodus","Leviticus","Numbers","Deuteronomy","Joshua","Judges","Ruth","1 Samuel","2 Samuel","1 Kings","2 Kings","1 Chronicles","2 Chronicles","Ezra","Nehemiah","Esther","Job","Psalms","Proverbs","Ecclesiastes","Song of Solomon","Isaiah","Jeremiah","Lamentations","Ezekiel","Daniel","Hosea","Joel","Amos","Obadiah","Jonah","Micah","Nahum","Habakkuk","Zephaniah","Haggai","Zechariah","Malachi","Matthew","Mark","Luke","John","Acts","Romans","1 Corinthians","2 Corinthians","Galatians","Ephesians","Philippians","Colossians","1 Thessalonians","2 Thessalonians","1 Timothy","2 Timothy","Titus","Philemon","Hebrews","James","1 Peter","2 Peter","1 John","2 John","3 John","Jude","Revelation"]
# inverse_index = {}
# # document_name = f"{book}_ch_{chapter}.docx"
# # document = Document(f"./Books/{book}/{document_name}")

# def make_name(term, max_len=4):
# 	query = term.lower().encode("UTF-8")
# 	return "{}.txt".format(hashlib.md5(query).hexdigest()[:max_len])

# def make_json(keyword, inverse_index, current_document, starting_index):
# 	max_ngram = 6
# 	for i in range(3, min(max_ngram, len(keyword)) + 1):
# 		ngram = keyword[:i]
# 		if (not ngram in inverse_index) and len(ngram) != 0:
# 			file_name = make_name(ngram.lower())
# 			inverse_index[ngram] = file_name
# 			txt_json = {
# 				current_document: [starting_index]
# 			}
# 			with open(f"./hashed_documents/{file_name}", "a") as file:
# 				file.write(f"\n{ngram}\t{json.dumps(txt_json)}")
# 		elif len(ngram) != 0:
# 			file_name = inverse_index[ngram]
# 			txt_json = {}
# 			all_lines = []
# 			with open(f"./hashed_documents/{file_name}", "r") as file:
# 				all_lines = file.readlines()

# 			found = False
# 			for index, line in enumerate(all_lines):
# 				if line.split("\t")[0].lower() == ngram.lower():
# 					found = True
# 					txt_json = json.loads(line.split("\t")[1])
# 					if not current_document in txt_json:
# 						txt_json[current_document] = [starting_index]
# 					else:
# 						txt_json[current_document].append(starting_index)
# 					all_lines[index] = f"{ngram}\t{json.dumps(txt_json)}"

# 			if len(ngram) != 0 and found == False:
# 				txt_json = {
# 					current_document: [starting_index]
# 				}
# 				all_lines.append(f"{ngram}\t{json.dumps(txt_json)}")

# 			with open(f"./hashed_documents/{file_name}", "w") as file:
# 				file.write("\n".join([line.strip() for line in all_lines]))

# def hash_book(book_name):
# 	num_of_chap = len(os.listdir(f"./Books/{book_name}")) - 1
# 	for i in range(1, num_of_chap):
# 		print("book: " + book_name + "\nchapter: " + str(i))
# 		document_name = f"{book_name}_ch_{i}.docx"
# 		document = Document(f"./Books/{book_name}/{document_name}")
# 		text = ""
# 		stop_words = {"a","about","above","after","again","against","all","am","an","and","any","are","aren't","as","at","be","because","been","before","being","below","between","both","but","by","can't","cannot","could","couldn't","did","didn't","do","does","doesn't","doing","don't","down","during","each","few","for","from","further","had","hadn't","has","hasn't","have","haven't","having","he","he'd","he'll","he's","her","here","here's","hers","herself","him","himself","his","how","how's","i","i'd","i'll","i'm","i've","if","in","into","is","isn't","it","it's","its","itself","let's","me","more","most","mustn't","my","myself","no","nor","not","of","off","on","once","only","or","other","ought","our","ours","ourselves","out","over","own","same","shan't","she","she'd","she'll","she's","should","shouldn't","so","some","such","than","that","that's","the","their","theirs","them","themselves","then","there","there's","these","they","they'd","they'll","they're","they've","this","those","through","to","too","under","until","up","very","was","wasn't","we","we'd","we'll","we're","we've","were","weren't","what","what's","when","when's","where","where's","which","while","who","who's","whom","why","why's","with","won't","would","wouldn't","you","you'd","you'll","you're","you've","your","yours","yourself","yourselves"}

# 		for paragraph in document.paragraphs:
# 			text = text + " " + paragraph.text
		
# 		for line_index, line in enumerate(text.split("\n")):
# 			for word_index, word in enumerate(line.split(" ")):
# 				if not word.lower() in stop_words and word != "":
# 					starting_index = 0
# 					for i in range(line_index):
# 						starting_index += len(text.split("\n")[i])
# 					starting_index += word_index
# 					word = word.strip(" \t\n.:[(<>)].,'\"‘’?--•…").lower()
# 					make_json(word, inverse_index, document_name, starting_index)

# def main():
# 	for book in books:
# 		hash_book(book)

# if __name__ == "__main__":
# 	main()




