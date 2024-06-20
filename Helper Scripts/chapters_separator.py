from docx import Document
import os

books = ["Genesis","Exodus","Leviticus","Numbers","Deuteronomy","Joshua","Judges","Ruth","1\\ Samuel","2\\ Samuel","1\\ Kings","2\\ Kings","1\\ Chronicles","2\\ Chronicles","Ezra","Nehemiah","Esther","Job","Psalms","Proverbs","Ecclesiastes","Song\\ of\\ Solomon","Isaiah","Jeremiah","Lamentations","Ezekiel","Daniel","Hosea","Joel","Amos","Obadiah","Jonah","Micah","Nahum","Habakkuk","Zephaniah","Haggai","Zechariah","Malachi","Matthew","Mark","Luke","John","Acts","Romans","1\\ Corinthians","2\\ Corinthians","Galatians","Ephesians","Philippians","Colossians","1\\ Thessalonians","2\\ Thessalonians","1\\ Timothy","2\\ Timothy","Titus","Philemon","Hebrews","James","1\\ Peter","2\\ Peter","1\\ John","2\\ John","3\\ John","Jude","Revelation"]
book_name = ["Genesis","Exodus","Leviticus","Numbers","Deuteronomy","Joshua","Judges","Ruth","1 Samuel","2 Samuel","1 Kings","2 Kings","1 Chronicles","2 Chronicles","Ezra","Nehemiah","Esther","Job","Psalms","Proverbs","Ecclesiastes","Song of Solomon","Isaiah","Jeremiah","Lamentations","Ezekiel","Daniel","Hosea","Joel","Amos","Obadiah","Jonah","Micah","Nahum","Habakkuk","Zephaniah","Haggai","Zechariah","Malachi","Matthew","Mark","Luke","John","Acts","Romans","1 Corinthians","2 Corinthians","Galatians","Ephesians","Philippians","Colossians","1 Thessalonians","2 Thessalonians","1 Timothy","2 Timothy","Titus","Philemon","Hebrews","James","1 Peter","2 Peter","1 John","2 John","3 John","Jude","Revelation"]

def validity_checker(query):
	if(query.lower() == book.lower() +" chapter"):
		return True
	return False

# def get_index(query):
# 	for i in range(len(books)):
# 		if (books[i].startswith(query[:1]) and books[i].endswith(query[len(books[i]) - 3:])):
# 			return i
# 	raise Exception("Sorry, no words match the book")

# for book in book_name:
# 	document = Document(f"./Books/{book}/{book}.docx")
# 	starting_index = 0
# 	chapter_number = 1
# 	for index, paragraph in enumerate(document.paragraphs):
# 		if(index < len(document.paragraphs) - 1 and index != 0 and validity_checker(paragraph.text.strip() + " " + document.paragraphs[index + 1].text.split(" ")[0].strip())):
# 			chapter = document.paragraphs[starting_index:index]
# 			starting_index = index
# 			file_name = f"rm ./Books/{books[get_index(book)]}/{chapter_number}.docx"
# 			os.system(file_name)
# 			chapter_number += 1
# 		if(index == len(document.paragraphs) - 1):
# 			chapter = document.paragraphs[starting_index:]
# 			starting_index = index
# 			file_name = f"rm ./Books/{books[get_index(book)]}/{chapter_number}.docx"
# 			os.system(file_name)


for book in book_name:
	document = Document(f"./Books/{book}/{book}.docx")
	starting_index = 0
	chapter_number = 1
	for index, paragraph in enumerate(document.paragraphs):
		if(index < len(document.paragraphs) - 1 and index != 0 and validity_checker(paragraph.text.strip() + " " + document.paragraphs[index + 1].text.split(" ")[0].strip())):
			chapter = document.paragraphs[starting_index:index]
			starting_index = index
			new_docx = Document()
			for paragraph in chapter:
				new_docx.add_paragraph(paragraph.text, style=paragraph.style)
			file_name = f"./Books/{book}/{book}_ch_{chapter_number}.docx"
			new_docx.save(file_name)
			chapter_number += 1
		if(index == len(document.paragraphs) - 1):
			chapter = document.paragraphs[starting_index:]
			starting_index = index
			new_docx = Document()
			for paragraph in chapter:
				new_docx.add_paragraph(paragraph.text, style=paragraph.style)
			file_name = f"./Books/{book}/{book}_ch_{chapter_number}.docx"
			new_docx.save(file_name)