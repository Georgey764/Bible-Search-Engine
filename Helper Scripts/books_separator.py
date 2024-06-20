from docx import Document

document = Document("bible.docx")

book = ["Genesis","Exodus","Leviticus","Numbers","Deuteronomy","Joshua","Judges","Ruth","1 Samuel","2 Samuel","1 Kings","2 Kings","1 Chronicles","2 Chronicles","Ezra","Nehemiah","Esther","Job","Psalms","Proverbs","Ecclesiastes","Song of Solomon","Isaiah","Jeremiah","Lamentations","Ezekiel","Daniel","Hosea","Joel","Amos","Obadiah","Jonah","Micah","Nahum","Habakkuk","Zephaniah","Haggai","Zechariah","Malachi","Matthew","Mark","Luke","John","Acts","Romans","1 Corinthians","2 Corinthians","Galatians","Ephesians","Philippians","Colossians","1 Thessalonians","2 Thessalonians","1 Timothy","2 Timothy","Titus","Philemon","Hebrews","James","1 Peter","2 Peter","1 John","2 John","3 John","Jude","Revelation"]
book_name = ["dummy", "Genesis","Exodus","Leviticus","Numbers","Deuteronomy","Joshua","Judges","Ruth","1 Samuel","2 Samuel","1 Kings","2 Kings","1 Chronicles","2 Chronicles","Ezra","Nehemiah","Esther","Job","Psalms","Proverbs","Ecclesiastes","Song of Solomon","Isaiah","Jeremiah","Lamentations","Ezekiel","Daniel","Hosea","Joel","Amos","Obadiah","Jonah","Micah","Nahum","Habakkuk","Zephaniah","Haggai","Zechariah","Malachi","Matthew","Mark","Luke","John","Acts","Romans","1 Corinthians","2 Corinthians","Galatians","Ephesians","Philippians","Colossians","1 Thessalonians","2 Thessalonians","1 Timothy","2 Timothy","Titus","Philemon","Hebrews","James","1 Peter","2 Peter","1 John","2 John","3 John","Jude","Revelation"]


def loop_over_name(query):
	for i in range(len(book)):
		if(query == book[i].lower() + "chapter 1"):
			return True
	return False

last_chapter_index = 0
chapters_made = 0
for index, paragraph in enumerate(document.paragraphs):
	if(index < len(document.paragraphs) - 1 and loop_over_name(paragraph.text.lower() + document.paragraphs[index + 1].text.lower())):
		chapters = document.paragraphs[last_chapter_index:index]
		print(chapters[0].text)
		new_docx = Document()
		for chapter in chapters:
			new_docx.add_paragraph(chapter.text, style=chapter.style)
		file_name = f"./Books/{book_name[chapters_made]}/{book_name[chapters_made]}.docx"
		chapters_made += 1
		last_chapter_index = index
		new_docx.save(file_name)

	if(index == len(document.paragraphs) - 1):
		chapters = document.paragraphs[last_chapter_index:]
		new_docx = Document()
		for chapter in chapters:
			new_docx.add_paragraph(chapter.text, style=chapter.style)
		file_name = f"./Books/{book_name[chapters_made]}/{book_name[chapters_made]}.docx"
		new_docx.save(file_name)